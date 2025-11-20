import os
import re
import io
import time
import threading
import traceback
import requests
from bs4 import BeautifulSoup
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
from pathlib import Path

# Загружаем .env при запуске
load_dotenv(Path(__file__).resolve().parent / ".env")

try:
    import tkinter as tk
    from tkinter import ttk, messagebox
except Exception:
    raise SystemExit("tkinter required")

# Настройки
GOOGLE_API_KEY = os.environ.get('GOOGLE_API_KEY')
GOOGLE_CX = os.environ.get('GOOGLE_CX')
GOOGLE_SEARCH_URL = 'https://www.googleapis.com/customsearch/v1'
MAX_RESULTS = 5
IMAGE_MAX_WIDTH_INCHES = 4
REQUEST_HEADERS = {"User-Agent": "Mozilla/5.0 (KAD-Bot/1.0)"}
MAX_COLLECTED_CHARS = 30000

# Вспомогательные функции
def exponential_backoff(max_attempts=5, initial_delay=1.0, factor=2.0):
    def deco(func):
        def wrapper(*args, **kwargs):
            delay = initial_delay
            attempt = 0
            while True:
                try:
                    return func(*args, **kwargs)
                except Exception as e:
                    attempt += 1
                    if attempt >= max_attempts:
                        raise
                    time.sleep(delay)
                    delay *= factor
        return wrapper
    return deco

def sanitize_filename(name: str) -> str:
    return re.sub(r'[\\/*:?"<>|]', '_', name.strip())

# Google Search
@exponential_backoff()
def google_search(query: str, num_results: int = MAX_RESULTS) -> list:
    if not GOOGLE_API_KEY or not GOOGLE_CX:
        raise RuntimeError("GOOGLE_API_KEY и GOOGLE_CX должны быть заданы")
    params = {'key': GOOGLE_API_KEY, 'cx': GOOGLE_CX, 'q': query, 'num': min(num_results, 10)}
    resp = requests.get(GOOGLE_SEARCH_URL, params=params, headers=REQUEST_HEADERS, timeout=15)
    resp.raise_for_status()
    data = resp.json()
    return [{'title': i.get('title', ''), 'snippet': i.get('snippet', ''), 'link': i.get('link', '')} for i in data.get('items', [])]

@exponential_backoff()
def google_search_image_url(query: str) -> str:
    params = {
        'key': GOOGLE_API_KEY,
        'cx': GOOGLE_CX,
        'q': query,
        'searchType': 'image',
        'num': 1,
        'imgSize': 'medium',  # Попробуем использовать фильтр по размеру изображения
    }
    resp = requests.get(GOOGLE_SEARCH_URL, params=params, headers=REQUEST_HEADERS, timeout=15)
    resp.raise_for_status()
    items = resp.json().get('items', [])
    return items[0].get('link', '') if items else ''

# Извлечение текста
@exponential_backoff()
def extract_text_from_url(url: str, max_chars: int = 4000) -> str:
    try:
        resp = requests.get(url, headers=REQUEST_HEADERS, timeout=12)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, 'html.parser')
        # Убираем ненужные теги
        for tag in soup(['script', 'style', 'nav', 'footer', 'header', 'form', 'aside', 'noscript']):
            tag.decompose()
        main = soup.find('main') or soup.find('article') or soup
        text = ' '.join(main.stripped_strings)
        return re.sub(r'\s+', ' ', text)[:max_chars]
    except Exception as e:
        return f"(Ошибка при загрузке {url}: {e})"

# Создание docx
def apply_paragraph_style(paragraph, text, bold=False, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.alignment = alignment
    return run

def build_docx(title: str, sections: list, sources: list, image_url: str, out_path: str):
    doc = Document()
    # Заголовок
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = 'Times New Roman'
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    # Содержание
    for heading, body in sections:
        h = doc.add_paragraph()
        apply_paragraph_style(h, heading, bold=True)
        b = doc.add_paragraph()
        apply_paragraph_style(b, body)
        doc.add_paragraph()

    # Источники
    s_head = doc.add_paragraph()
    apply_paragraph_style(s_head, 'Источники:', bold=True)
    for src in sources:
        p = doc.add_paragraph()
        apply_paragraph_style(p, src)
    
    # Изображение
    if image_url:
        try:
            # Попробуем скачать изображение
            img_data = requests.get(image_url, headers=REQUEST_HEADERS, timeout=15).content
            # Проверка на валидность изображения
            try:
                img = Image.open(io.BytesIO(img_data))
                img.verify()  # Проверяем, что изображение корректное
                img_stream = io.BytesIO(img_data)

                # Сохраняем изображение во временный файл
                image_filename = sanitize_filename(title[:40]) + ".jpg"
                image_path = os.path.join(os.getcwd(), image_filename)
                with open(image_path, "wb") as f:
                    f.write(img_data)

                # Вставляем изображение в документ
                doc.add_page_break()
                pic = doc.add_picture(img_stream, width=Inches(IMAGE_MAX_WIDTH_INCHES))
                doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            except Exception as img_err:
                doc.add_paragraph(f"(Ошибка при обработке изображения: {img_err})")

        except Exception as e:
            doc.add_paragraph(f"(Не удалось загрузить изображение: {e})")

    doc.save(out_path)

# Основная логика
def execute_ai_plan(prompt: str, log_fn=print) -> str:
    log_fn("Поиск информации в интернете...")
    sections = []
    sources = []
    image_url = ''
    
    # Определяем шаблон разделов в зависимости от типа объекта
    character_keywords = ["мультфильм", "аниме", "сериал", "игра", "персонаж", "герой"]
    if any(word.lower() in prompt.lower() for word in character_keywords):
        queries = {
            "Описание": f"{prompt} персонаж описание",
            "Характер": f"{prompt} персонаж черты характера",
            "Сюжет/Роль": f"{prompt} персонаж сюжет роль",
            "Интересные факты": f"{prompt} персонаж интересные факты"
        }
    else:
        # Шаблон для реальных людей
        queries = {
            "Биография (рождение и детство, дата рождения, смерть)": f"Биография {prompt}",
            "Деятельность": f"Деятельность {prompt}",
            "Труды/работы": f"Труды {prompt}",
            "Вклад в историю": f"Вклад в историю {prompt}"
        }

    for heading, query in queries.items():
        try:
            results = google_search(query)
            collected_text = ""
            for r in results:
                url = r['link']
                text = extract_text_from_url(url)
                collected_text += text + "\n\n"
                sources.append(url)
                log_fn(f" → {heading}: загружен текст с {url}")
                time.sleep(0.2)
            sections.append((heading, collected_text.strip()))
        except Exception as e:
            log_fn(f"Ошибка поиска '{query}': {e}")
            sections.append((heading, ''))

    # Поиск изображения
    try:
        image_url = google_search_image_url(f"фото {prompt}")
        log_fn(f"Найдено изображение: {image_url}")
    except Exception as e:
        log_fn(f"Ошибка поиска изображения: {e}")

    # Введение
    intro_text = f"Статья по теме: {prompt}\nИнформация собрана из открытых источников."
    sections.insert(0, ('Главная', intro_text))

    # Создать docx
    filename = sanitize_filename(prompt[:40]) or 'article'
    out_path = os.path.join(os.getcwd(), f"{filename}.docx")
    build_docx(prompt, sections, list(dict.fromkeys(sources)), image_url, out_path)
    log_fn(f"Готово! Файл: {out_path}")
    return out_path

# GUI
class KADApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Когнитивный Ассемблер Документов (КАД)")
        self.geometry("800x560")
        frm = ttk.Frame(self, padding=12)
        frm.pack(fill=tk.BOTH, expand=True)
        ttk.Label(frm, text='Тема статьи:').pack(anchor=tk.W)
        self.topic_var = tk.StringVar()
        ttk.Entry(frm, textvariable=self.topic_var).pack(fill=tk.X, pady=6)
        self.generate_btn = ttk.Button(frm, text='Сгенерировать', command=self.on_generate)
        self.generate_btn.pack(pady=6)
        ttk.Label(frm, text='Лог:').pack(anchor=tk.W)
        self.log_box = tk.Text(frm, height=25)
        self.log_box.pack(fill=tk.BOTH, expand=True)
        ttk.Button(frm, text='Открыть папку', command=self.open_cwd).pack(pady=6)

    def log(self, *args):
        self.log_box.insert(tk.END, ' '.join(str(a) for a in args)+'\n')
        self.log_box.see(tk.END)
        self.update_idletasks()

    def on_generate(self):
        topic = self.topic_var.get().strip()
        if not topic:
            messagebox.showwarning('Пустая тема', 'Введите тему')
            return
        self.generate_btn.config(state=tk.DISABLED)
        threading.Thread(target=self.run_generate, args=(topic,), daemon=True).start()

    def run_generate(self, topic):
        try:
            out = execute_ai_plan(topic, log_fn=self.log)
            messagebox.showinfo('Готово', f'Файл сохранён: {out}')
        except Exception as e:
            tb = traceback.format_exc()
            self.log('Ошибка:', e, '\n', tb)
            messagebox.showerror('Ошибка', str(e))
        finally:
            self.generate_btn.config(state=tk.NORMAL)

    def open_cwd(self):
        import subprocess, sys
        path = os.getcwd()
        if sys.platform == 'win32':
            os.startfile(path)
        elif sys.platform == 'darwin':
            subprocess.Popen(['open', path])
        else:
            subprocess.Popen(['xdg-open', path])

if __name__ == "__main__":
    app = KADApp()
    app.mainloop()